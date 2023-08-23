# Build by Tony Esposito
import streamlit as st
from azure.cognitiveservices.vision.computervision import ComputerVisionClient
from msrest.authentication import CognitiveServicesCredentials
from PIL import Image
from docx import Document  # Import for Word document
from docx.shared import Inches  # Import for setting image size
import os
import openai  # Add OpenAI
import time
from io import BytesIO

# Function to generate a well-formatted step-by-step Word document
def generate_word_document(text, image_path, doc_path):
    doc = Document()
    doc.add_heading('Step-by-Step Procedure', 0)
    doc.add_paragraph('The following document provides a step-by-step procedure based on extracted text and image.')
    steps = text.split('. ')
    for i, step in enumerate(steps):
        if step:
            doc.add_heading(f'Step {i + 1}', level=1)
            doc.add_paragraph(step)
    doc.add_heading('Screenshot:', level=1)
    doc.add_picture(image_path, width=Inches(4.0))
    doc.save(doc_path)

# Load keys from environment variables for better security
subscription_key = os.getenv("AZURE_SUBSCRIPTION_KEY", "xxxxxxxxxx")
endpoint = os.getenv("AZURE_ENDPOINT", "https://xxxxx.cognitiveservices.azure.com/")
openai_api_key = os.getenv("OPENAI_API_KEY", "sk-xxxxxxxxxx")  # Add OpenAI API Key

# Initialize Azure Computer Vision client
computervision_client = ComputerVisionClient(endpoint, CognitiveServicesCredentials(subscription_key))
# Initialize OpenAI API key
openai.api_key = openai_api_key  # Add OpenAI initialization

st.title('Extract Text from Image and Generate Word Document')

uploaded_file = st.file_uploader('Upload Image', type=['png', 'jpeg', 'jpg'])

if uploaded_file is not None:
    image = Image.open(uploaded_file)
    byte_stream = BytesIO()
    image.save(byte_stream, format="PNG")
    byte_stream.seek(0)
    
    # Save the image for adding it to the Word document later
    image_path = 'uploaded_image.png'
    image.save(image_path)

    raw_response = computervision_client.read_in_stream(byte_stream, raw=True)
    operation_location = raw_response.headers["Operation-Location"]
    operation_id = operation_location.split("/")[-1]

    while True:
        read_result = computervision_client.get_read_result(operation_id)
        if read_result.status.lower() not in ['notstarted', 'running']:
            break
        time.sleep(1)

    if read_result.status.lower() == 'succeeded':
        text_azure = ""
        for result in read_result.analyze_result.read_results:
            for line in result.lines:
                text_azure += line.text + " "

        st.write('Azure Computer Vision Extracted Text:')
        st.write(text_azure)

        # Generate the Word document
        doc_path = 'Step_by_Step_Procedure.docx'
        generate_word_document(text_azure, image_path, doc_path)

        # Add download button for the Word document
        with open(doc_path, 'rb') as file:
            st.download_button(
                label="Download Word Document",
                data=file,
                file_name="Step_by_Step_Procedure.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        st.success(f'Generated {doc_path} with extracted texts from Azure')

else:
    st.write('Upload an image to extract text using ACV and generate a Word document. @by Tony Esposito')
